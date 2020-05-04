
import subprocess
import os
import re
import pymysql
import docx
import jieba
db = pymysql.connect("localhost", "root", "", "专利法律判决书库", charset='utf8')
cursor = db.cursor()
sql = "select * from `专利法律判决书库`.`allinfo` "
cursor.execute(sql)
results = cursor.fetchall()
# print(results)

# 写入败诉案件
# loseCount = 0
# for row in results:
#     money_all = 0
#     print(row[0])
#     file = docx.Document(row[1])
#     for p in file.paragraphs:
#         if not bool(re.search("辩称", p.text)):
#             if not bool(re.search("诉请", p.text)):
#                 if not bool(re.search("请求法院", p.text)):
#                     if bool(re.search("驳回原告", p.text)) and bool(re.search("诉讼请求", p.text)):
#                         print(p.text)
#                         money_all = 0
#                         print(money_all)
#                         sql = """UPDATE `专利法律判决书库`.`allinfo` SET `赔偿金额` = %s WHERE `id` =  %s"""
#                         cursor.execute(sql, (money_all, row[0]))
#                         loseCount = loseCount + 1
#                         break
#
#     # subprocess.call("pause", shell=True)
# print("败诉的案件数量："+str(loseCount))

# # 写入胜诉案件
# winCount = 0
# panjueruxai_count = 0
# jingjisunshi_count = 0
# un_jingjisunshi_count = 0
# for row in results:
#     money_all = 0
#     print(row[0])
#     # print(type(row[10]))
#     file = docx.Document(row[1])
#     if row[10] is None :
#         winCount = winCount + 1
#         fullText = ''
#         for p in file.paragraphs:
#             fullText = fullText + p.text
#         # print(fullText)
#         if bool(re.search("判决如下", fullText)):
#             panjueruxia = fullText.split("判决如下")[1]
#             # print(panjueruxia)
#             panjueruxai_count = panjueruxai_count + 1
#             if bool(re.search("经济损失", panjueruxia)):
#                 jingjisunshi = panjueruxia.split("经济损失")[1]
#                 print(jingjisunshi)
#                 jingjisunshi_count = jingjisunshi_count + 1
#             else:
#                 un_jingjisunshi_count = un_jingjisunshi_count + 1
#                 print(panjueruxia)
#
#
#     # subprocess.call("pause", shell=True)
# print("胜诉的案件数量："+str(winCount))
# print("判决如下数量："+str(panjueruxai_count))
# print("经济损失数量："+str(jingjisunshi_count))
# print("非经济损失数量："+str(un_jingjisunshi_count))

