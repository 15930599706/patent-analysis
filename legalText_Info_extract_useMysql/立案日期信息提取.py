import subprocess
import os
import re
import pymysql
import docx
import jieba
db = pymysql.connect("localhost", "root", "", "专利法律判决书库", charset='utf8')
cursor = db.cursor()
sql = "select * from `专利法律判决书库`.`allinfo` "
cursor.execute(sql)
results = cursor.fetchall()
# print(results)

# 识别判决书名称是否规范
# nonstandard_txt_count = 0
# for row in results:
#     file = docx.Document(row[1])
#     # print(row[2])
#     # 原告和被告信息
#     Plaintiff_Defendant = re.split("侵害",row[2])[0]
#     # print(Plaintiff_Defendant)
#     if not bool(re.search("与", Plaintiff_Defendant)):
#         # print(Plaintiff_Defendant)
#         # print(row[1])
#         sql = """UPDATE `专利法律判决书库`.`allinfo` SET `name是否规范` = %s WHERE `id` =  %s"""
#         cursor.execute(sql,('no', row[0]))
#
#         nonstandard_txt_count = nonstandard_txt_count + 1
#     else:
#         sql = """UPDATE `专利法律判决书库`.`allinfo` SET `name是否规范` = %s WHERE `id` =  %s"""
#         cursor.execute(sql, ('yes', row[0]))
# print("不规范文件的总数为：",nonstandard_txt_count)



# 对名称规范的判决书，提取原被告信息
# for row in results:
#     # file = docx.Document(row[1])
#     # print(row[2])
#     if row[3] == 'yes':
#         # 原告和被告信息
#         Plaintiff_Defendant = re.split("侵害",row[2])[0]
#         Plaintiff = re.split("与", Plaintiff_Defendant)[0]
#         Defendant = re.split("与", Plaintiff_Defendant)[1]
#         # print(Plaintiff_Defendant)
#         # print(Plaintiff)
#         sql = """UPDATE `专利法律判决书库`.`allinfo` SET `原告` = %s WHERE `id` =  %s"""
#         cursor.execute(sql,(Plaintiff, row[0]))
#         # print(Defendant)
#         sql = """UPDATE `专利法律判决书库`.`allinfo` SET `被告` = %s WHERE `id` =  %s"""
#         cursor.execute(sql,(Defendant, row[0]))
#


# 对名称不规范的判决书，提取原被告信息
for row in results:
    Plaintiff = ''
    Defendant = ''
    if row[3] == 'no':
        # 处理二审判决
        if  bool(re.search("二审", row[2])):
            print("当前二审判决书为：《《《《", row[2])
            rowCount = 0
            file = docx.Document(row[1])
            for p in file.paragraphs:
                # 原被告信息在前3-20行
                if rowCount >= 4:
                    # print(p.text)
                    # print(p.text.startswith('　　上诉人'))
                    if  bool(p.text.startswith("　　上诉人（")) :
                        tmp_str = re.split('：',p.text)[1]
                        tmp_str2 = re.split('，|。', tmp_str)[0]
                        Plaintiff = Plaintiff + '、'+ tmp_str2
                        print("原告是->" + tmp_str2)

                    if bool(p.text.startswith("　　被上诉人（")) :
                        tmp_str = re.split('：', p.text)[1]
                        tmp_str2 = re.split('，|。', tmp_str)[0]
                        Defendant = Defendant + '、' + tmp_str2
                        print("被告是->" + tmp_str2)

                if rowCount >= 20:
                    break
                rowCount = rowCount + 1

            if Plaintiff == '':
                finallPlaintiff = '张三罗'
            else:
                finallPlaintiff = Plaintiff[1:]
            print("所有的原告是=>" + finallPlaintiff)
            if Defendant == '':
                finallDefendant = '张三罗'
            else:
                finallDefendant = Defendant[1:]
            print("所有的被告是=>" + finallDefendant)

            sql = """UPDATE `专利法律判决书库`.`allinfo` SET `原告` = %s WHERE `id` =  %s"""
            cursor.execute(sql,(finallPlaintiff, row[0]))
            sql = """UPDATE `专利法律判决书库`.`allinfo` SET `被告` = %s WHERE `id` =  %s"""
            cursor.execute(sql,(finallDefendant, row[0]))
            # subprocess.call("pause", shell=True)

        # 处理一审判决
        else:
            print("当前一审判决书为：《《《《", row[2])
            rowCount = 0
            file = docx.Document(row[1])
            for p in file.paragraphs:
                # 原被告信息在前3-20行
                if rowCount >= 4:
                    # print(p.text)
                    # print(p.text.startswith('　　原告：'))
                    if bool(p.text.startswith("　　原告：")):
                        ano_tmp_str = re.split('：', p.text)[1]
                        ano_tmp_str2 = re.split('，|。', ano_tmp_str)[0]
                        Plaintiff = Plaintiff + '、'+ ano_tmp_str2
                        # print("原告是->" + ano_tmp_str2)

                    if bool(p.text.startswith("　　被告：")):
                        ano_tmp_str = re.split('：', p.text)[1]
                        ano_tmp_str2 = re.split('，|。', ano_tmp_str)[0]
                        Defendant = Defendant + '、' + ano_tmp_str2
                        # print("被告是->" + ano_tmp_str2)

                if rowCount >= 20:
                    break
                rowCount = rowCount + 1

            if Plaintiff == '':
                finallPlaintiff = '张三罗'
            else:
                finallPlaintiff = Plaintiff[1:]
            print("所有的原告是=>" + finallPlaintiff)
            if Defendant == '':
                finallDefendant = '张三罗'
            else:
                finallDefendant = Defendant[1:]
            print("所有的被告是=>" + finallDefendant)

            sql = """UPDATE `专利法律判决书库`.`allinfo` SET `原告` = %s WHERE `id` =  %s"""
            cursor.execute(sql,(finallPlaintiff, row[0]))
            sql = """UPDATE `专利法律判决书库`.`allinfo` SET `被告` = %s WHERE `id` =  %s"""
            cursor.execute(sql,(finallDefendant, row[0]))
            # subprocess.call("pause", shell=True)

